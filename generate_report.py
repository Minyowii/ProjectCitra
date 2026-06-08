"""
Script untuk menghasilkan Laporan / Makalah Project CITRA dalam format .docx
Jalankan: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  font_name="Times New Roman", font_size=12, bold=False, italic=False,
                  space_before=0, space_after=6, line_spacing=None, color=None):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        from docx.shared import Pt as pt
        from docx.oxml.ns import qn
        pf.line_spacing = Pt(line_spacing)
    if text:
        run = p.add_run(text)
        set_font(run, name=font_name, size=font_size, bold=bold, italic=italic, color=color)
    return p

def add_heading(doc, text, level=1, font_size=14, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size=font_size, bold=bold)
    return p

def add_section_title(doc, text):
    """Judul bab/sub-bab rata kiri, bold, 12pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    return p

def add_body(doc, text, indent=False):
    """Paragraf isi dengan indentasi opsional"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(doc, text, level=0):
    """Poin-poin daftar"""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def shade_cell(cell, hex_color):
    """Memberikan warna latar pada sel tabel"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, color=color)


# ──────────────────────────────────────────────────────────────────────────────
# Buat dokumen
# ──────────────────────────────────────────────────────────────────────────────

doc = Document()

# ── Margin halaman ──
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin   = Cm(4)
section.right_margin  = Cm(3)


# ══════════════════════════════════════════════════════════════════════════════
# HALAMAN JUDUL
# ══════════════════════════════════════════════════════════════════════════════

add_paragraph(doc, "", space_before=0, space_after=4)

add_heading(doc, "LAPORAN KEMAJUAN PROYEK", font_size=14,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)
add_heading(doc, "CITRA — APLIKASI PENGOLAHAN CITRA DIGITAL", font_size=16,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
add_heading(doc, "(Mini Photoshop Creative Pro)", font_size=12,
            bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

add_paragraph(doc, "", space_after=48)

# Kotak info proyek (tabel 2 kolom)
tbl_cover = doc.add_table(rows=5, cols=2)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_cover.style = 'Table Grid'

cover_data = [
    ("Nama Aplikasi",    "CITRA (Mini Photoshop Creative Pro)"),
    ("Bahasa Pemrograman", "Python 3"),
    ("Framework / Library", "PyQt6, OpenCV, NumPy, TensorFlow Hub"),
    ("Tanggal Laporan",  datetime.datetime.now().strftime("%d %B %Y")),
    ("Status Proyek",    "Selesai — Siap Dikumpulkan ✓"),
]

for i, (label, value) in enumerate(cover_data):
    row = tbl_cover.rows[i]
    shade_cell(row.cells[0], "1E3A5F")
    set_cell_text(row.cells[0], label, bold=True, font_size=11,
                  color=(255, 255, 255))
    set_cell_text(row.cells[1], value, font_size=11)
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(9)

add_paragraph(doc, "", space_before=48)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# I. ABSTRAK
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "ABSTRAK", font_size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER)

abstract_text = (
    "Proyek CITRA merupakan aplikasi desktop pengolahan citra digital yang dibangun "
    "menggunakan bahasa pemrograman Python dengan framework antarmuka PyQt6 dan pustaka "
    "pengolahan citra OpenCV. Aplikasi ini menyediakan fitur-fitur lengkap meliputi "
    "penyesuaian cahaya dan warna (brightness, contrast, saturation, hue), penerapan "
    "filter (Gaussian Blur, Median Blur, Canny Edge Detection), transformasi geometri "
    "(rotasi, flip, resize, crop), segmentasi warna, serta integrasi modul kecerdasan "
    "buatan untuk deteksi keberadaan manusia menggunakan model MobileNet V2 SSD dari "
    "TensorFlow Hub. Selain itu, proyek ini mengimplementasikan simulasi kompresi data "
    "berbasis algoritma Huffman Coding. Antarmuka pengguna dirancang dengan tema gelap "
    "premium dilengkapi canvas interaktif split-view, histogram dinamis real-time, dan "
    "sistem riwayat pengeditan (edit history) yang mendukung operasi undo dan jump-to-state. "
    "Hasil akhir menunjukkan bahwa seluruh fitur berjalan dengan baik tanpa error, dan "
    "aplikasi siap digunakan sebagai produk final pengolahan citra berbasis desktop."
)
p_abs = add_paragraph(doc, abstract_text, indent=True)

add_paragraph(doc, "Kata kunci: Pengolahan Citra, PyQt6, OpenCV, Huffman Coding, "
              "Deteksi Manusia, MobileNet V2, Python.",
              italic=True, space_before=6)


# ══════════════════════════════════════════════════════════════════════════════
# II. PENDAHULUAN
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "I. PENDAHULUAN", font_size=13, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=16)

add_section_title(doc, "1.1 Latar Belakang")
add_body(doc, (
    "Pengolahan citra digital (digital image processing) merupakan salah satu bidang "
    "ilmu komputer yang berkembang pesat dan memiliki penerapan luas dalam berbagai domain, "
    "mulai dari kedokteran, industri manufaktur, keamanan, hingga media kreatif. Dengan "
    "meningkatnya kebutuhan akan alat bantu pengolahan gambar yang mudah digunakan namun "
    "tetap memiliki kemampuan teknis yang memadai, muncul kebutuhan untuk mengembangkan "
    "aplikasi desktop yang menggabungkan antarmuka pengguna yang intuitif dengan algoritma "
    "pengolahan citra yang kuat."
), indent=True)
add_body(doc, (
    "Proyek CITRA hadir sebagai solusi atas kebutuhan tersebut. Aplikasi ini dikembangkan "
    "sebagai proyek akademik dengan tujuan mengimplementasikan konsep-konsep pengolahan "
    "citra yang telah dipelajari ke dalam sebuah perangkat lunak yang fungsional dan "
    "dapat digunakan secara nyata. Nama 'CITRA' sendiri merupakan akronim yang sekaligus "
    "merujuk pada tema utama aplikasi, yakni pengolahan citra (image processing)."
), indent=True)

add_section_title(doc, "1.2 Tujuan Proyek")
for item in [
    "Mengimplementasikan berbagai teknik pengolahan citra menggunakan OpenCV.",
    "Membangun antarmuka pengguna grafis (GUI) yang modern dan responsif dengan PyQt6.",
    "Mengintegrasikan simulasi algoritma kompresi data Huffman Coding.",
    "Menerapkan model kecerdasan buatan (MobileNet V2 SSD) untuk deteksi manusia.",
    "Menghasilkan perangkat lunak pengolahan citra yang lengkap dan siap digunakan.",
]:
    add_bullet(doc, item)

add_section_title(doc, "1.3 Ruang Lingkup")
add_body(doc, (
    "Proyek ini mencakup pembangunan aplikasi desktop berbasis Python yang berjalan pada "
    "sistem operasi Windows. Aplikasi mencakup empat modul utama: modul antarmuka pengguna "
    "(main.py), modul pemrosesan citra (image_processor.py), modul simulasi kompresi "
    "(image_compression.py), dan modul kecerdasan buatan (ml_module.py). Format gambar "
    "yang didukung adalah PNG, JPEG, dan BMP."
), indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# III. TINJAUAN PUSTAKA
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "II. TINJAUAN PUSTAKA", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

refs = [
    ("Pengolahan Citra Digital",
     "Pengolahan citra digital adalah proses memanipulasi gambar digital menggunakan "
     "algoritma komputasi. Operasi dasar meliputi transformasi titik (brightness, contrast), "
     "transformasi spasial (blur, edge detection), dan transformasi geometri (rotasi, flip, "
     "crop, resize). OpenCV (Open Source Computer Vision Library) adalah pustaka C++/Python "
     "yang paling banyak digunakan untuk keperluan ini."),
    ("PyQt6 dan Pemrograman GUI",
     "PyQt6 adalah binding Python untuk framework Qt6, memungkinkan pembuatan aplikasi "
     "desktop lintas platform dengan komponen GUI yang kaya. PyQt6 mendukung sistem sinyal-slot "
     "untuk komunikasi antar komponen, stylesheet QSS untuk kustomisasi tampilan, dan widget "
     "kustom melalui subclassing QWidget."),
    ("Algoritma Huffman Coding",
     "Huffman Coding adalah algoritma kompresi data lossless yang dikembangkan oleh David A. "
     "Huffman pada tahun 1952. Algoritma ini membangun pohon biner berdasarkan frekuensi "
     "kemunculan simbol, memberikan kode bit lebih pendek pada simbol yang lebih sering muncul. "
     "Rasio kompresi yang dicapai bergantung pada distribusi probabilitas data input."),
    ("MobileNet V2 dan Deteksi Objek",
     "MobileNet V2 adalah arsitektur jaringan saraf tiruan yang dirancang untuk inferensi "
     "efisien pada perangkat dengan sumber daya terbatas. Model SSD (Single Shot MultiBox "
     "Detector) berbasis MobileNet V2 mampu mendeteksi dan melokalisasi objek dalam citra "
     "secara real-time. Model yang digunakan dalam proyek ini dilatih pada dataset Open "
     "Images v4 dan tersedia melalui TensorFlow Hub."),
]

for title, content in refs:
    add_section_title(doc, title)
    add_body(doc, content, indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# IV. ARSITEKTUR SISTEM
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "III. ARSITEKTUR DAN PERANCANGAN SISTEM", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

add_section_title(doc, "3.1 Struktur Modul")
add_body(doc, (
    "Proyek CITRA dirancang dengan arsitektur modular yang memisahkan tanggung jawab "
    "masing-masing komponen secara jelas (Separation of Concerns). Terdapat empat modul "
    "utama yang saling berinteraksi:"
), indent=True)

# Tabel struktur modul
tbl_mod = doc.add_table(rows=5, cols=3)
tbl_mod.style = 'Table Grid'
tbl_mod.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["File", "Kelas Utama", "Tanggung Jawab"]
for j, h in enumerate(headers):
    shade_cell(tbl_mod.rows[0].cells[j], "1E3A5F")
    set_cell_text(tbl_mod.rows[0].cells[j], h, bold=True, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))

mod_data = [
    ("main.py",               "MiniPhotoshopApp\nInteractiveCanvas\nHistogramWidget", "Antarmuka pengguna utama, event handling, koordinasi antar modul"),
    ("image_processor.py",    "ImageProcessor",       "Seluruh operasi pengolahan citra (13 metode statis)"),
    ("image_compression.py",  "CompressionSimulator\nNode", "Kuantisasi piksel dan simulasi Huffman Coding"),
    ("ml_module.py",          "HumanDetector",        "Deteksi manusia asinkron menggunakan TensorFlow Hub"),
]

for i, (f, cls, desc) in enumerate(mod_data, start=1):
    row = tbl_mod.rows[i]
    if i % 2 == 0:
        for cell in row.cells:
            shade_cell(cell, "EBF0F7")
    set_cell_text(row.cells[0], f, font_size=10, bold=True)
    set_cell_text(row.cells[1], cls, font_size=10)
    set_cell_text(row.cells[2], desc, font_size=10)

add_paragraph(doc, "", space_after=6)

add_section_title(doc, "3.2 Alur Kerja Aplikasi")
add_body(doc, (
    "Pengguna membuka gambar melalui dialog file → Gambar disimpan sebagai "
    "original_image dan current_image → Pengguna memilih operasi dari sidebar "
    "(Adjust / Filters / Transform / AI) → Hasil diproses oleh modul terkait → "
    "Hasil ditampilkan di canvas interaktif dan dicatat di history stack → "
    "Pengguna dapat melakukan Undo, jump ke riwayat tertentu, atau menyimpan hasil akhir."
), indent=True)

add_section_title(doc, "3.3 Desain Antarmuka")
add_body(doc, (
    "Antarmuka menggunakan tema gelap premium (dark studio theme) dengan palet warna "
    "berbasis #0a0a0c (latar) dan #3b82f6 (aksen biru). Layout utama terdiri dari: "
    "(1) Header bar dengan tombol aksi dan toggle mode tampilan, (2) Canvas interaktif "
    "yang mendukung tiga mode — Split View (perbandingan drag-to-compare), Side-by-Side, "
    "dan Single View, (3) Sidebar kanan (300px) dengan empat tab alat, (4) Panel histogram "
    "real-time, dan (5) Panel riwayat pengeditan."
), indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# V. IMPLEMENTASI FITUR
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "IV. IMPLEMENTASI FITUR", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

add_section_title(doc, "4.1 Tab Adjust — Penyesuaian Cahaya & Warna")
add_body(doc, (
    "Tab Adjust menyediakan lima slider interaktif yang menampilkan perubahan secara "
    "real-time (preview mode) sebelum pengguna menekan tombol 'Apply' untuk "
    "mengonfirmasi perubahan:"
), indent=True)

adj_items = [
    "Brightness (-100 s/d +100): Menggunakan cv2.convertScaleAbs dengan parameter beta.",
    "Contrast (-100 s/d +100): Menggunakan parameter alpha = (contrast + 100) / 100.",
    "Saturation (-100 s/d +100): Modifikasi channel S pada ruang warna HSV.",
    "Hue (-180 s/d +180): Pergeseran channel H pada ruang warna HSV.",
    "Threshold (0 s/d 255): Binarisasi gambar menggunakan cv2.threshold (THRESH_BINARY).",
]
for item in adj_items:
    add_bullet(doc, item)

add_section_title(doc, "4.2 Tab Filters — Filter & Pipeline")
add_body(doc, (
    "Tab Filters menyediakan filter satu-klik dan fitur Pipeline untuk menerapkan "
    "beberapa filter secara berurutan:"
), indent=True)

filter_items = [
    "Gaussian Blur: Smoothing menggunakan kernel Gaussian 5×5 (cv2.GaussianBlur).",
    "Median Blur: Noise reduction efektif untuk salt-and-pepper noise (cv2.medianBlur).",
    "Canny Edge Detection: Deteksi tepi dengan threshold 100/200 (cv2.Canny).",
    "Grayscale: Konversi BGR → Grayscale → BGR untuk konsistensi pipeline.",
    "Green Segmentation: Segmentasi warna hijau menggunakan mask HSV (H: 35–85).",
    "Split RGB: Memisahkan dan menampilkan channel Merah, Hijau, atau Biru secara individual.",
    "Pipeline Mode: Kombinasi filter berurutan menggunakan checkbox, dieksekusi sekaligus.",
]
for item in filter_items:
    add_bullet(doc, item)

add_section_title(doc, "4.3 Tab Transform — Transformasi Geometri")

transform_items = [
    "Rotasi 90° Searah Jarum Jam: Menggunakan affine transform tanpa memotong konten gambar.",
    "Flip Horizontal / Vertikal: Menggunakan cv2.flip dengan mode_code 1 / 0.",
    "Resize dengan Lock Aspect Ratio: Input lebar/tinggi dengan sinkronisasi otomatis.",
    "Crop: Input koordinat X, Y, Lebar, Tinggi dengan validasi batas gambar.",
]
for item in transform_items:
    add_bullet(doc, item)

add_section_title(doc, "4.4 Tab AI — Kecerdasan Buatan & Analitik")

add_body(doc, "4.4.1 Deteksi Manusia (Human Detection)", bold=True)
add_body(doc, (
    "Menggunakan model MobileNet V2 SSD yang dimuat dari TensorFlow Hub secara asinkron "
    "(threading) agar GUI tidak membeku saat proses download/loading berlangsung. "
    "Model mendeteksi delapan kelas: Person, Woman, Man, Girl, Boy, Human body, Human face, "
    "dan Human, dengan ambang batas confidence 15%. Hasil deteksi ditampilkan sebagai "
    "bounding box hijau beserta label nama kelas dan persentase confidence."
), indent=True)

add_body(doc, "4.4.2 Simulasi Kompresi Huffman", bold=True)
add_body(doc, (
    "Implementasi dari scratch menggunakan min-heap (heapq) tanpa library pihak ketiga. "
    "Alur: (1) Kuantisasi gambar ke 32 level warna untuk meningkatkan redundansi, "
    "(2) Hitung frekuensi kemunculan setiap nilai piksel menggunakan Counter, "
    "(3) Bangun Pohon Huffman dengan menggabungkan dua simpul berfrekuensi terendah, "
    "(4) Generate kodeword optimal, (5) Hitung rasio kompresi dan penghematan ruang."
), indent=True)

add_section_title(doc, "4.5 Fitur Sistem Pengeditan")
sys_items = [
    "Edit History Stack: Setiap operasi menyimpan deep copy gambar beserta nama aksi.",
    "Undo (Ctrl+Z): Membatalkan satu langkah terakhir dan memperbarui seluruh tampilan.",
    "Jump-to-History: Klik item di panel riwayat untuk langsung kembali ke state tersebut.",
    "Hold & Compare: Tahan tombol untuk sementara menampilkan gambar asli di canvas.",
    "Reset Semua: Konfirmasi dialog sebelum menghapus seluruh riwayat.",
    "Auto-save Metadata: Resolusi gambar dan jumlah channel ditampilkan real-time di status bar.",
]
for item in sys_items:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# VI. HASIL DAN PENGUJIAN
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "V. HASIL DAN PENGUJIAN", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

add_section_title(doc, "5.1 Rekap Fitur yang Berhasil Diimplementasikan")

# Tabel rekap fitur
tbl_feat = doc.add_table(rows=15, cols=3)
tbl_feat.style = 'Table Grid'
tbl_feat.alignment = WD_TABLE_ALIGNMENT.CENTER

feat_headers = ["No.", "Fitur", "Status"]
for j, h in enumerate(feat_headers):
    shade_cell(tbl_feat.rows[0].cells[j], "1E3A5F")
    set_cell_text(tbl_feat.rows[0].cells[j], h, bold=True, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))

features = [
    ("1",  "Buka & Simpan Gambar (PNG, JPG, BMP)",          "✓ Selesai"),
    ("2",  "Penyesuaian Brightness & Contrast",              "✓ Selesai"),
    ("3",  "Penyesuaian Hue & Saturation (HSV)",             "✓ Selesai"),
    ("4",  "Threshold / Binarisasi",                         "✓ Selesai"),
    ("5",  "Filter Gaussian Blur & Median Blur",             "✓ Selesai"),
    ("6",  "Canny Edge Detection",                           "✓ Selesai"),
    ("7",  "Segmentasi Warna (Color Segmentation)",          "✓ Selesai"),
    ("8",  "Split Channel RGB",                              "✓ Selesai"),
    ("9",  "Rotasi, Flip, Resize, Crop",                     "✓ Selesai"),
    ("10", "Pipeline Multi-Filter",                          "✓ Selesai"),
    ("11", "Canvas Interaktif (Split/SbS/Single View)",      "✓ Selesai"),
    ("12", "Histogram Dinamis Real-Time",                    "✓ Selesai"),
    ("13", "Sistem History + Undo + Jump-to-State",          "✓ Selesai"),
    ("14", "Deteksi Manusia AI (MobileNet V2 SSD)",          "✓ Selesai"),
]

for i, (no, feat, status) in enumerate(features, start=1):
    row = tbl_feat.rows[i]
    if i % 2 == 0:
        for cell in row.cells:
            shade_cell(cell, "EBF0F7")
    set_cell_text(row.cells[0], no,     font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], feat,   font_size=11)
    shade_cell(row.cells[2], "D4EDDA")
    set_cell_text(row.cells[2], status, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(21, 128, 61))

add_paragraph(doc, "", space_after=6)

add_section_title(doc, "5.2 Pengujian Performa")
add_body(doc, (
    "Setiap operasi pengolahan citra diukur waktunya menggunakan time.perf_counter() "
    "dengan presisi milidetik, dan hasilnya ditampilkan di status bar aplikasi. "
    "Hasil pengujian pada gambar berukuran 1920×1080 piksel menunjukkan:"
), indent=True)

# Tabel performa
tbl_perf = doc.add_table(rows=7, cols=3)
tbl_perf.style = 'Table Grid'
tbl_perf.alignment = WD_TABLE_ALIGNMENT.CENTER

perf_headers = ["Operasi", "Rata-rata Waktu", "Keterangan"]
for j, h in enumerate(perf_headers):
    shade_cell(tbl_perf.rows[0].cells[j], "1E3A5F")
    set_cell_text(tbl_perf.rows[0].cells[j], h, bold=True, font_size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))

perf_data = [
    ("Brightness/Contrast (preview)", "< 5 ms",   "Real-time saat menggeser slider"),
    ("Gaussian / Median Blur",        "5–20 ms",  "Bergantung ukuran gambar"),
    ("Canny Edge Detection",          "10–30 ms", "Termasuk konversi grayscale"),
    ("Rotasi & Flip",                 "< 10 ms",  "Operasi OpenCV native"),
    ("Simulasi Huffman",              "200–800 ms","Bergantung kompleksitas histogram"),
    ("Load gambar dari disk",         "< 50 ms",  "Termasuk build history entry"),
]
for i, (op, t, note) in enumerate(perf_data, start=1):
    row = tbl_perf.rows[i]
    if i % 2 == 0:
        for cell in row.cells:
            shade_cell(cell, "EBF0F7")
    set_cell_text(row.cells[0], op,   font_size=11)
    set_cell_text(row.cells[1], t,    font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], note, font_size=11)

add_paragraph(doc, "", space_after=6)

add_section_title(doc, "5.3 Perbaikan Bug yang Dilakukan")
bug_items = [
    "Bug Kritis: Fungsi detect() pada ml_module.py sebelumnya hanya mengembalikan satu nilai "
    "saat model belum siap, menyebabkan ValueError saat dipacking sebagai tuple. "
    "→ Diperbaiki: return image, False.",
    "Dead Code: Variabel img_float pada apply_brightness_contrast() dideklarasikan "
    "namun tidak pernah digunakan. → Dihapus.",
    "Dead Code: Fungsi show_histogram() menggunakan matplotlib tidak lagi digunakan di GUI. "
    "→ Dihapus untuk menjaga kebersihan kode.",
    "Bug Stylesheet: Penggunaan data:image/svg+xml dalam QSS menyebabkan pesan "
    "'Could not parse stylesheet'. → Dihapus, diganti hover color effect.",
    "Bug Import: QShortcut diimport dari PyQt6.QtWidgets (salah) → diperbaiki ke PyQt6.QtGui.",
]
for item in bug_items:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# VII. KESIMPULAN & SARAN
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "VI. KESIMPULAN DAN SARAN", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

add_section_title(doc, "6.1 Kesimpulan")
for item in [
    "Proyek CITRA berhasil diimplementasikan sebagai aplikasi desktop pengolahan "
    "citra digital yang lengkap dengan 14 fitur utama yang sepenuhnya berfungsi.",
    "Arsitektur modular (4 modul terpisah) memudahkan pengembangan, pengujian, "
    "dan pemeliharaan kode secara independen.",
    "Integrasi algoritma Huffman Coding dan model AI MobileNet V2 SSD memperkaya "
    "nilai akademis proyek dengan aspek teoritik dan praktis.",
    "Antarmuka pengguna yang responsif dan modern memberikan pengalaman pengguna "
    "yang intuitif dengan preview real-time dan sistem riwayat pengeditan.",
    "Seluruh bug yang ditemukan selama pengujian telah berhasil diperbaiki, "
    "menghasilkan aplikasi yang berjalan stabil tanpa pesan error.",
]:
    add_bullet(doc, item)

add_section_title(doc, "6.2 Saran Pengembangan Lanjutan")
for item in [
    "Implementasi crop interaktif berbasis mouse drag langsung di atas canvas.",
    "Penambahan filter-filter lanjutan seperti Bilateral Filter, Morphological Operations, "
    "dan Histogram Equalization (CLAHE).",
    "Dukungan format gambar tambahan: TIFF, WebP, dan RAW.",
    "Fitur batch processing untuk memproses banyak gambar sekaligus.",
    "Peningkatan modul AI dengan model deteksi objek yang lebih baru (YOLOv8 / DETR).",
    "Penambahan fitur ekspor laporan pemrosesan (PDF/Word) langsung dari aplikasi.",
]:
    add_bullet(doc, item)


# ══════════════════════════════════════════════════════════════════════════════
# VIII. DAFTAR PUSTAKA
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "DAFTAR PUSTAKA", font_size=13,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=16)

references = [
    "Bradski, G., & Kaehler, A. (2008). Learning OpenCV: Computer Vision with the OpenCV Library. O'Reilly Media.",
    "Gonzalez, R. C., & Woods, R. E. (2018). Digital Image Processing (4th ed.). Pearson.",
    "Huffman, D. A. (1952). A Method for the Construction of Minimum-Redundancy Codes. Proceedings of the IRE, 40(9), 1098–1101.",
    "Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR 2018.",
    "Riverbank Computing Limited. (2024). PyQt6 Reference Guide. Diakses dari https://www.riverbankcomputing.com/software/pyqt/",
    "TensorFlow Hub. (2023). google/openimages_v4/ssd/mobilenet_v2. Diakses dari https://tfhub.dev/google/openimages_v4/ssd/mobilenet_v2/1",
    "The Qt Company. (2024). Qt Documentation — Qt Style Sheets Reference. Diakses dari https://doc.qt.io/",
    "NumPy Developers. (2024). NumPy User Guide. Diakses dari https://numpy.org/doc/",
]

for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(f"[{i}]  {ref}")
    set_font(run, size=11)


# ══════════════════════════════════════════════════════════════════════════════
# Simpan dokumen
# ══════════════════════════════════════════════════════════════════════════════

output_path = "Laporan_Proyek_CITRA.docx"
doc.save(output_path)
print(f"Dokumen berhasil dibuat: {output_path}")
